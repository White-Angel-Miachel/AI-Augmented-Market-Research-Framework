"""
Report Parser Module
Handles parsing of various document formats (PDF, DOCX, TXT)
"""

from pathlib import Path
from typing import Optional
import re


class ReportParser:
    """
    Parser for industry reports in various formats.
    Supports PDF, DOCX, and TXT files.
    """
    
    def __init__(self, chunk_size: int = 4000, overlap: int = 200):
        """
        Initialize the parser.
        
        Args:
            chunk_size: Maximum characters per chunk for processing
            overlap: Character overlap between chunks
        """
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def parse(self, file_path: str) -> str:
        """
        Parse a report file and extract text content.
        
        Args:
            file_path: Path to the report file
            
        Returns:
            Extracted text content
        """
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        if suffix == '.pdf':
            return self._parse_pdf(path)
        elif suffix == '.docx':
            return self._parse_docx(path)
        elif suffix == '.txt':
            return self._parse_txt(path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    def _parse_pdf(self, path: Path) -> str:
        """Parse PDF file."""
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("PyPDF2 required for PDF parsing. Install with: pip install PyPDF2")
        
        reader = PdfReader(str(path))
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return self._clean_text('\n\n'.join(text_parts))
    
    def _parse_docx(self, path: Path) -> str:
        """Parse Word document."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required for DOCX parsing. Install with: pip install python-docx")
        
        doc = Document(str(path))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return self._clean_text('\n\n'.join(text_parts))
    
    def _parse_txt(self, path: Path) -> str:
        """Parse plain text file."""
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return self._clean_text(f.read())
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove page numbers and headers/footers patterns
        text = re.sub(r'\n\d+\n', '\n', text)
        text = re.sub(r'Page \d+ of \d+', '', text)
        
        return text.strip()
    
    def chunk_text(self, text: str) -> list:
        """
        Split text into overlapping chunks for processing.
        
        Args:
            text: Full text content
            
        Returns:
            List of text chunks
        """
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence end
                for punct in ['. ', '.\n', '! ', '? ']:
                    last_punct = text[start:end].rfind(punct)
                    if last_punct > self.chunk_size * 0.5:
                        end = start + last_punct + 1
                        break
            
            chunks.append(text[start:end])
            start = end - self.overlap
        
        return chunks
    
    def extract_sections(self, text: str) -> dict:
        """
        Attempt to extract named sections from the report.
        
        Args:
            text: Report text content
            
        Returns:
            Dictionary mapping section names to content
        """
        sections = {}
        
        # Common section headers
        header_patterns = [
            r'^#+\s*(.+)$',  # Markdown headers
            r'^([A-Z][A-Z\s]+)$',  # ALL CAPS headers
            r'^\d+\.?\s+([A-Z][^.]+)$',  # Numbered sections
        ]
        
        current_section = "Introduction"
        current_content = []
        
        for line in text.split('\n'):
            is_header = False
            
            for pattern in header_patterns:
                match = re.match(pattern, line.strip())
                if match and len(match.group(1)) < 100:
                    # Save previous section
                    if current_content:
                        sections[current_section] = '\n'.join(current_content)
                    
                    current_section = match.group(1).strip()
                    current_content = []
                    is_header = True
                    break
            
            if not is_header:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
