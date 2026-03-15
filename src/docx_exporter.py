"""
DOCX Exporter for Market Research Analysis Reports
Converts Markdown analysis reports to professionally formatted Word documents.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


class DocxExporter:
    """Converts Markdown analysis reports to formatted DOCX documents."""

    # Color scheme
    PRIMARY = RGBColor(0x1A, 0x3C, 0x6E)       # Deep navy blue
    SECONDARY = RGBColor(0x2E, 0x86, 0xAB)     # Teal accent
    TEXT_DARK = RGBColor(0x2D, 0x2D, 0x2D)      # Near-black for body
    TEXT_LIGHT = RGBColor(0x5A, 0x5A, 0x5A)     # Grey for metadata
    ACCENT = RGBColor(0xE8, 0x6C, 0x00)         # Orange accent

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles for a professional look."""
        # Page margins
        for section in self.doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.color.rgb = self.TEXT_DARK
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

        # Title style
        title_style = self.doc.styles['Title']
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = self.PRIMARY
        title_style.paragraph_format.space_after = Pt(4)

        # Heading 1
        h1 = self.doc.styles['Heading 1']
        h1.font.name = 'Calibri'
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = self.PRIMARY
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(8)

        # Heading 2
        h2 = self.doc.styles['Heading 2']
        h2.font.name = 'Calibri'
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = self.SECONDARY
        h2.paragraph_format.space_before = Pt(14)
        h2.paragraph_format.space_after = Pt(6)

        # Heading 3
        h3 = self.doc.styles['Heading 3']
        h3.font.name = 'Calibri'
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.color.rgb = self.PRIMARY
        h3.paragraph_format.space_before = Pt(10)
        h3.paragraph_format.space_after = Pt(4)

    def _add_horizontal_rule(self):
        """Add a visual separator line."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # Add a bottom border to the paragraph
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single',
            qn('w:sz'): '6',
            qn('w:space'): '1',
            qn('w:color'): '2E86AB'
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_metadata_line(self, label: str, value: str):
        """Add a styled metadata line (e.g., Sector: Enterprise)."""
        p = self.doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_label.font.color.rgb = self.TEXT_LIGHT
        run_value = p.add_run(value)
        run_value.font.size = Pt(11)
        run_value.font.color.rgb = self.TEXT_DARK

    def _add_bullet_point(self, text: str, level: int = 0):
        """Add a formatted bullet point with bold keys."""
        p = self.doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
        p.paragraph_format.space_after = Pt(3)

        # Parse bold markers **text**
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = self.PRIMARY
            else:
                run = p.add_run(part)
                run.font.size = Pt(10.5)
                run.font.color.rgb = self.TEXT_DARK

    def _add_paragraph_with_bold(self, text: str):
        """Add a paragraph that handles **bold** markers inline."""
        p = self.doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

    def _add_table(self, rows: list):
        """Add a styled table to the document. First row is treated as header."""
        if not rows:
            return
        num_cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'

        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx >= num_cols:
                    break
                cell = table.cell(row_idx, col_idx)
                cell.text = ''
                p = cell.paragraphs[0]
                # Clean bold markers
                clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                run = p.add_run(clean_text)
                run.font.size = Pt(9.5)
                run.font.name = 'Calibri'

                if row_idx == 0:
                    # Header row styling
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    # Navy background
                    shading_elm = cell._element.get_or_add_tcPr()
                    shading = shading_elm.makeelement(qn('w:shd'), {
                        qn('w:fill'): '1A3C6E',
                        qn('w:val'): 'clear'
                    })
                    shading_elm.append(shading)
                else:
                    run.font.color.rgb = self.TEXT_DARK
                    # Alternating row shading
                    if row_idx % 2 == 0:
                        shading_elm = cell._element.get_or_add_tcPr()
                        shading = shading_elm.makeelement(qn('w:shd'), {
                            qn('w:fill'): 'EBF5FB',
                            qn('w:val'): 'clear'
                        })
                        shading_elm.append(shading)

        # Add spacing after table
        self.doc.add_paragraph()

    def convert_md_to_docx(self, md_path: str, docx_path: str = None) -> str:
        """
        Convert a Markdown analysis report to a formatted DOCX document.

        Args:
            md_path: Path to the source Markdown file
            docx_path: Optional output path (defaults to same name with .docx)

        Returns:
            Path to the generated DOCX file
        """
        md_path = Path(md_path)
        if docx_path is None:
            docx_path = md_path.with_suffix('.docx')
        else:
            docx_path = Path(docx_path)

        # Read markdown content
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Reset document
        self.doc = Document()
        self._setup_styles()

        # Parse and convert
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            # Skip empty lines
            if not line:
                i += 1
                continue

            # Title (# )
            if line.startswith('# ') and not line.startswith('## '):
                title_text = line[2:].strip()
                # Clean up filename-style titles
                title_text = title_text.replace('_', ' ').replace('-', ' ')
                p = self.doc.add_paragraph(title_text, style='Title')
                self._add_horizontal_rule()
                i += 1
                continue

            # Heading 2 (## )
            if line.startswith('## ') and not line.startswith('### '):
                heading_text = line[3:].strip()
                self.doc.add_heading(heading_text, level=1)
                i += 1
                continue

            # Heading 3 (### )
            if line.startswith('### '):
                heading_text = line[4:].strip()
                self.doc.add_heading(heading_text, level=2)
                i += 1
                continue

            # Horizontal rule (---)
            if line.strip() == '---':
                self._add_horizontal_rule()
                i += 1
                continue

            # Heading 4 (#### )
            if line.startswith('#### '):
                heading_text = line[5:].strip()
                p = self.doc.add_heading(heading_text, level=3)
                i += 1
                continue

            # Markdown table (| col | col |)
            if line.strip().startswith('|') and '|' in line[1:]:
                table_rows = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    row_line = lines[i].strip()
                    # Skip separator rows (|---|---|)
                    if re.match(r'^\|[\s\-:|]+\|$', row_line):
                        i += 1
                        continue
                    cells = [c.strip() for c in row_line.split('|')[1:-1]]
                    table_rows.append(cells)
                    i += 1
                if table_rows:
                    self._add_table(table_rows)
                continue

            # Blockquote (> text)
            if line.startswith('> '):
                quote_text = line[2:].strip()
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                # Add left border
                pPr = p._element.get_or_add_pPr()
                pBdr = pPr.makeelement(qn('w:pBdr'), {})
                left = pBdr.makeelement(qn('w:left'), {
                    qn('w:val'): 'single',
                    qn('w:sz'): '12',
                    qn('w:space'): '4',
                    qn('w:color'): '2E86AB'
                })
                pBdr.append(left)
                pPr.append(pBdr)
                # Parse bold in blockquote
                parts = re.split(r'(\*\*.*?\*\*)', quote_text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                        run.font.color.rgb = self.SECONDARY
                    else:
                        run = p.add_run(part)
                        run.font.color.rgb = self.TEXT_LIGHT
                i += 1
                continue

            # Metadata lines (**Key:** Value)
            meta_match = re.match(r'^\*\*(.+?):\*\*\s*(.*)', line)
            if meta_match:
                self._add_metadata_line(meta_match.group(1), meta_match.group(2))
                i += 1
                continue

            # Sub-bullets (*   **text**)
            if line.strip().startswith('*   ') or line.strip().startswith('*  '):
                bullet_text = re.sub(r'^\s*\*\s+', '', line)
                level = 1 if line.startswith('    ') or line.startswith('\t') else 0
                self._add_bullet_point(bullet_text, level)
                i += 1
                continue

            # Numbered list items (1. , 2. , etc.)
            num_match = re.match(r'^\s*(\d+)\.\s+(.*)', line)
            if num_match:
                text = num_match.group(2)
                p = self.doc.add_paragraph(style='List Number')
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                        run.font.color.rgb = self.PRIMARY
                    else:
                        p.add_run(part)
                i += 1
                continue

            # Bullet points (- **Key:** Value)
            if line.startswith('- '):
                bullet_text = line[2:]
                self._add_bullet_point(bullet_text, level=0)
                i += 1
                continue

            # Code blocks (```json or ```)  - render as plain text
            if line.strip().startswith('```'):
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1)
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = self.TEXT_LIGHT
                i += 1  # skip closing ```
                continue

            # Regular paragraph (with bold support)
            self._add_paragraph_with_bold(line)
            i += 1

        # Save
        self.doc.save(str(docx_path))
        return str(docx_path)
